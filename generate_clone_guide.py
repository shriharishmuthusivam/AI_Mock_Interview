from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

GREEN = RGBColor(0x1F, 0x4E, 0x38)
GRAY = RGBColor(0x40, 0x40, 0x40)


def add_code_block(code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x0B, 0x53, 0x94)


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    return table


# Title
title = doc.add_heading("AI Mock Interviewer - Clone & Setup Guide", level=0)
for run in title.runs:
    run.font.color.rgb = GREEN
    run.font.size = Pt(24)

p = doc.add_paragraph()
r = p.add_run("How to clone this project on a new device, including all requirements and commands.")
r.italic = True
r.font.color.rgb = GRAY

# 1. Overview
doc.add_heading("1. Project Overview", level=1)
doc.add_paragraph(
    "AI Mock Interviewer is a web application that lets students take AI-driven mock interviews "
    "and lets interviewers create syllabi, upload students, and generate PDF reports. "
    "The project has two parts:"
)
add_table(
    ["Folder", "What it is"],
    [
        ["backend/", "Express + MongoDB API server (runs on port 5000)"],
        ["frontend/", "React app created with Create React App (dev server on port 3000)"],
    ],
)

# 2. Requirements
doc.add_heading("2. Requirements", level=1)
doc.add_paragraph("Install the following on the new device before cloning:")

add_bullet("v2 or newer - https://git-scm.com/downloads", "Git ")
add_bullet(
    "LTS version 18 or 20+ - https://nodejs.org/en/download. (npm is included with Node.js.)",
    "Node.js ",
)
add_bullet(
    "Community Server (recommended) from https://www.mongodb.com/try/download/community, "
    "or a free MongoDB Atlas cluster at https://www.mongodb.com/atlas.",
    "MongoDB ",
)
add_bullet(
    "API key (free tier is enough) from https://console.groq.com. Interviews will not work without it.",
    "Groq ",
)
add_bullet(
    "a Gmail / any SMTP account if you want PDF report emails. Optional.",
    "Optional: ",
)

doc.add_paragraph()
doc.add_paragraph("Verify the installs with:")
add_code_block("git --version\nnode --version\nnpm --version")

# 3. Clone
doc.add_heading("3. Clone the Repository", level=1)
add_code_block("git clone https://github.com/shriharishmuthusivam/AI_Mock_Interview.git\ncd AI_Mock_Interview")

# 4. MongoDB setup
doc.add_heading("4. MongoDB Setup (Local, Recommended)", level=1)
add_bullet("Download MongoDB Community Server (Windows MSI or Linux package).")
add_bullet("Install it - on Windows keep the defaults so MongoDB runs as a Windows service (auto-start).")
add_bullet("Verify it is running:")
add_code_block("# Windows\nGet-Service MongoDB        # should show 'Running'\n\n# Linux\nsudo systemctl status mongod  # should show 'active (running)'")
add_bullet(
    "No database creation is needed - the app creates the ai_mock_interview database automatically on first run.",
    "Note: ",
)
add_bullet("MongoDB Compass (https://www.mongodb.com/products/compass) can be installed to browse the data.")

# 5. Backend
doc.add_heading("5. Backend Setup", level=1)
add_code_block("cd backend\nnpm install")
doc.add_paragraph("Create the environment file:")
add_code_block("copy .env.example .env")
doc.add_paragraph("Fill in backend/.env with real values:")
add_table(
    ["Variable", "Required", "What it is"],
    [
        ["GROQ_API_KEY", "Yes", "AI key from https://console.groq.com"],
        ["MONGO_URI", "Yes", "Local default: mongodb://127.0.0.1:27017/ai_mock_interview. Atlas example: mongodb+srv://<user>:<password>@<cluster>.mongodb.net/ai_mock_interview?retryWrites=true&w=majority"],
        ["JWT_SECRET", "Yes", "Any long random string (change it)"],
        ["CLIENT_URL", "No", "Frontend origin allowed by CORS (default http://localhost:3000)"],
        ["SMTP_HOST / SMTP_PORT / SMTP_SECURE / SMTP_USER / SMTP_PASS / SMTP_FROM", "No", "Report email. If unset, reports are generated but email is skipped."],
    ],
)
doc.add_paragraph("Run the server:")
add_code_block("npm start")
doc.add_paragraph("Or, for auto-reload while developing:")
add_code_block("npm run dev")
doc.add_paragraph("You should see: MongoDB Connected and Server running on port 5000.")

# 6. Frontend
doc.add_heading("6. Frontend Setup", level=1)
add_code_block("cd frontend\nnpm install\nnpm start")
doc.add_paragraph(
    "Opens the app at http://localhost:3000. No configuration is needed in development - "
    "the frontend already points at http://localhost:5000 and CORS defaults to http://localhost:3000.",
)

# 7. First-time app setup
doc.add_heading("7. First-Time App Setup", level=1)
add_bullet("Register an interviewer account (this also stores their email for PDF reports).")
add_bullet("Log in as the interviewer, open Setup, add a syllabus for each class, and set the question count.")
add_bullet("In Setup, upload students using the provided Excel template.")
add_bullet("Log in as a student and take an interview to confirm the AI (Groq) works.")

# 8. Known gotchas
doc.add_heading("8. Known Gotchas & Troubleshooting", level=1)
add_bullet(
    "Class names are hardcoded in two places and must stay in sync: "
    "frontend/src/constants.js (CLASSES) and backend/server.js (CLASSES). "
    "If the college uses different class names, update both.",
)
add_bullet("The Groq API key is mandatory - no interviews without it.")
add_bullet("Each developer's database is independent - there is no sync between local databases.")
add_bullet(
    "Camera/mic and live video only work on localhost or HTTPS (browser security). "
    "A plain http://<ip> setup will block the camera and the Jitsi live room.",
)
add_bullet(
    "Known unfixed bug: the dashboard/PDF report Feedback field can include the AI's "
    "'Next Question: ...' text (parseReply in backend/server.js).",
)
add_bullet(
    "The completed flag used by the interviewer dashboard is set when a report is generated; "
    "existing data can be backfilled once with: node backend/scripts/backfill-completed.js",
)

# 9. Production deployment note
doc.add_heading("9. Production Deployment (Optional)", level=1)
doc.add_paragraph(
    "A render.yaml is included in the repository root for deploying the backend to Render. "
    "When deploying, set the same environment variables (MONGO_URI, GROQ_API_KEY, JWT_SECRET, "
    "CLIENT_URL, and SMTP settings) in the Render dashboard, and deploy the frontend build "
    "(frontend/build) to a static host such as Netlify.",
)

doc.add_paragraph()
note = doc.add_paragraph()
nr = note.add_run(
    "Handoff note: when sharing the project folder, delete backend/node_modules, frontend/node_modules, "
    "frontend/build, backend/.env, and any *.log files first. The root .gitignore already excludes "
    "node_modules, .env, build, and *.log."
)
nr.italic = True
nr.font.size = Pt(10)

doc.save("AI_Mock_Interview_Clone_Guide.docx")
print("Saved AI_Mock_Interview_Clone_Guide.docx")
